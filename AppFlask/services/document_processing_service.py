import os
import io
import json
import logging
from datetime import datetime
from typing import Optional, Dict, List, Tuple
from werkzeug.utils import secure_filename
from PIL import Image, ImageDraw, ImageFont
import pytesseract
import cv2
import numpy as np
from pdf2image import convert_from_path, convert_from_bytes
import fitz  # PyMuPDF
from docx import Document
try:
    import magic
    MAGIC_AVAILABLE = True
except ImportError:
    MAGIC_AVAILABLE = False
import base64
from AppFlask.db import db_connection
from psycopg2.extras import RealDictCursor

logger = logging.getLogger(__name__)

class DocumentProcessingService:
    """Service de traitement avancé des documents"""
    
    def __init__(self):
        self.upload_folder = "uploads"
        self.preview_folder = "uploads/previews"
        self.thumbnail_folder = "uploads/thumbnails"
        self.ocr_folder = "uploads/ocr"
        
        # Créer les dossiers s'ils n'existent pas
        for folder in [self.preview_folder, self.thumbnail_folder, self.ocr_folder]:
            os.makedirs(folder, exist_ok=True)
        
        # Configuration Tesseract (adapter le chemin selon l'installation)
        if os.name == 'nt':  # Windows
            pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'
    
    def _get_mime_type_from_extension(self, file_path: str) -> str:
        """Obtenir le type MIME depuis l'extension de fichier (fallback si magic n'est pas disponible)"""
        ext = os.path.splitext(file_path)[1].lower()
        mime_types = {
            '.pdf': 'application/pdf',
            '.doc': 'application/msword',
            '.docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            '.jpg': 'image/jpeg',
            '.jpeg': 'image/jpeg',
            '.png': 'image/png',
            '.gif': 'image/gif',
            '.bmp': 'image/bmp',
            '.tiff': 'image/tiff',
            '.txt': 'text/plain',
            '.html': 'text/html',
            '.xml': 'text/xml',
            '.csv': 'text/csv'
        }
        return mime_types.get(ext, 'application/octet-stream')
    
    def extract_metadata(self, file_path: str) -> Dict:
        """Extraire les métadonnées d'un document"""
        try:
            metadata = {
                'file_size': os.path.getsize(file_path),
                'last_modified': datetime.fromtimestamp(os.path.getmtime(file_path)).isoformat(),
                'mime_type': magic.from_file(file_path, mime=True) if MAGIC_AVAILABLE else self._get_mime_type_from_extension(file_path),
                'pages': 0,
                'text_content': '',
                'language': 'fr',
                'has_images': False,
                'word_count': 0,
                'character_count': 0
            }
            
            file_ext = os.path.splitext(file_path)[1].lower()
            
            if file_ext == '.pdf':
                metadata.update(self._extract_pdf_metadata(file_path))
            elif file_ext in ['.docx', '.doc']:
                metadata.update(self._extract_docx_metadata(file_path))
            elif file_ext in ['.jpg', '.jpeg', '.png', '.tiff', '.bmp']:
                metadata.update(self._extract_image_metadata(file_path))
            
            return metadata
            
        except Exception as e:
            logger.error(f"Erreur lors de l'extraction des métadonnées: {str(e)}")
            return {'error': str(e)}
    
    def _extract_pdf_metadata(self, file_path: str) -> Dict:
        """Extraire les métadonnées d'un PDF"""
        try:
            doc = fitz.open(file_path)
            metadata = {
                'pages': doc.page_count,
                'title': doc.metadata.get('title', ''),
                'author': doc.metadata.get('author', ''),
                'subject': doc.metadata.get('subject', ''),
                'creator': doc.metadata.get('creator', ''),
                'producer': doc.metadata.get('producer', ''),
                'creation_date': doc.metadata.get('creationDate', ''),
                'modification_date': doc.metadata.get('modDate', ''),
                'has_images': False,
                'text_content': ''
            }
            
            # Extraire le texte de toutes les pages
            full_text = ""
            has_images = False
            
            for page_num in range(doc.page_count):
                page = doc[page_num]
                text = page.get_text()
                full_text += text + "\n"
                
                # Vérifier s'il y a des images
                if page.get_images():
                    has_images = True
            
            metadata['text_content'] = full_text.strip()
            metadata['has_images'] = has_images
            metadata['word_count'] = len(full_text.split())
            metadata['character_count'] = len(full_text)
            
            doc.close()
            return metadata
            
        except Exception as e:
            logger.error(f"Erreur lors de l'extraction des métadonnées PDF: {str(e)}")
            return {'error': str(e)}
    
    def _extract_docx_metadata(self, file_path: str) -> Dict:
        """Extraire les métadonnées d'un document Word"""
        try:
            doc = Document(file_path)
            
            # Extraire le texte
            full_text = ""
            for paragraph in doc.paragraphs:
                full_text += paragraph.text + "\n"
            
            metadata = {
                'pages': 1,  # Approximation
                'text_content': full_text.strip(),
                'word_count': len(full_text.split()),
                'character_count': len(full_text),
                'has_images': len(doc.inline_shapes) > 0,
                'paragraphs': len(doc.paragraphs),
                'tables': len(doc.tables)
            }
            
            # Métadonnées du document
            core_props = doc.core_properties
            if core_props.title:
                metadata['title'] = core_props.title
            if core_props.author:
                metadata['author'] = core_props.author
            if core_props.subject:
                metadata['subject'] = core_props.subject
            if core_props.created:
                metadata['creation_date'] = core_props.created.isoformat()
            if core_props.modified:
                metadata['modification_date'] = core_props.modified.isoformat()
            
            return metadata
            
        except Exception as e:
            logger.error(f"Erreur lors de l'extraction des métadonnées DOCX: {str(e)}")
            return {'error': str(e)}
    
    def _extract_image_metadata(self, file_path: str) -> Dict:
        """Extraire les métadonnées d'une image"""
        try:
            with Image.open(file_path) as img:
                metadata = {
                    'width': img.width,
                    'height': img.height,
                    'format': img.format,
                    'mode': img.mode,
                    'has_images': True
                }
                
                # Informations EXIF
                if hasattr(img, '_getexif') and img._getexif():
                    exif = img._getexif()
                    if exif:
                        metadata['exif'] = {}
                        for tag_id, value in exif.items():
                            tag = Image.ExifTags.TAGS.get(tag_id, tag_id)
                            metadata['exif'][tag] = value
                
                return metadata
                
        except Exception as e:
            logger.error(f"Erreur lors de l'extraction des métadonnées image: {str(e)}")
            return {'error': str(e)}
    
    def generate_thumbnail(self, file_path: str, size: Tuple[int, int] = (200, 200)) -> Optional[str]:
        """Générer une miniature du document"""
        try:
            file_name = os.path.basename(file_path)
            name_without_ext = os.path.splitext(file_name)[0]
            thumbnail_path = os.path.join(self.thumbnail_folder, f"{name_without_ext}_thumb.png")
            
            file_ext = os.path.splitext(file_path)[1].lower()
            
            if file_ext == '.pdf':
                return self._generate_pdf_thumbnail(file_path, thumbnail_path, size)
            elif file_ext in ['.jpg', '.jpeg', '.png', '.tiff', '.bmp']:
                return self._generate_image_thumbnail(file_path, thumbnail_path, size)
            elif file_ext in ['.docx', '.doc']:
                return self._generate_document_thumbnail(file_path, thumbnail_path, size)
            else:
                # Créer une miniature générique
                return self._generate_generic_thumbnail(file_name, thumbnail_path, size)
                
        except Exception as e:
            logger.error(f"Erreur lors de la génération de la miniature: {str(e)}")
            return None
    
    def _generate_pdf_thumbnail(self, file_path: str, thumbnail_path: str, size: Tuple[int, int]) -> str:
        """Générer une miniature à partir d'un PDF"""
        try:
            # Convertir la première page en image
            pages = convert_from_path(file_path, first_page=1, last_page=1, dpi=150)
            if pages:
                page = pages[0]
                page.thumbnail(size, Image.Resampling.LANCZOS)
                page.save(thumbnail_path, "PNG")
                return thumbnail_path
        except Exception as e:
            logger.error(f"Erreur PDF thumbnail: {str(e)}")
            return self._generate_generic_thumbnail("PDF", thumbnail_path, size)
    
    def _generate_image_thumbnail(self, file_path: str, thumbnail_path: str, size: Tuple[int, int]) -> str:
        """Générer une miniature à partir d'une image"""
        try:
            with Image.open(file_path) as img:
                img.thumbnail(size, Image.Resampling.LANCZOS)
                img.save(thumbnail_path, "PNG")
                return thumbnail_path
        except Exception as e:
            logger.error(f"Erreur image thumbnail: {str(e)}")
            return self._generate_generic_thumbnail("IMG", thumbnail_path, size)
    
    def _generate_document_thumbnail(self, file_path: str, thumbnail_path: str, size: Tuple[int, int]) -> str:
        """Générer une miniature pour un document Word"""
        # Pour l'instant, créer une miniature générique
        return self._generate_generic_thumbnail("DOC", thumbnail_path, size)
    
    def _generate_generic_thumbnail(self, file_type: str, thumbnail_path: str, size: Tuple[int, int]) -> str:
        """Créer une miniature générique avec le type de fichier"""
        try:
            img = Image.new('RGB', size, color='#2D3748')
            draw = ImageDraw.Draw(img)
            
            # Essayer de charger une police, sinon utiliser la police par défaut
            try:
                font = ImageFont.truetype("arial.ttf", 20)
            except:
                font = ImageFont.load_default()
            
            # Centrer le texte
            text = file_type
            bbox = draw.textbbox((0, 0), text, font=font)
            text_width = bbox[2] - bbox[0]
            text_height = bbox[3] - bbox[1]
            x = (size[0] - text_width) // 2
            y = (size[1] - text_height) // 2
            
            draw.text((x, y), text, fill='white', font=font)
            img.save(thumbnail_path, "PNG")
            return thumbnail_path
            
        except Exception as e:
            logger.error(f"Erreur generic thumbnail: {str(e)}")
            return None
    
    def generate_preview(self, file_path: str, page: int = 1) -> Optional[str]:
        """Générer un aperçu du document"""
        try:
            file_name = os.path.basename(file_path)
            name_without_ext = os.path.splitext(file_name)[0]
            preview_path = os.path.join(self.preview_folder, f"{name_without_ext}_preview_p{page}.png")
            
            file_ext = os.path.splitext(file_path)[1].lower()
            
            if file_ext == '.pdf':
                return self._generate_pdf_preview(file_path, preview_path, page)
            elif file_ext in ['.jpg', '.jpeg', '.png', '.tiff', '.bmp']:
                return self._generate_image_preview(file_path, preview_path)
            else:
                return None
                
        except Exception as e:
            logger.error(f"Erreur lors de la génération de l'aperçu: {str(e)}")
            return None
    
    def _generate_pdf_preview(self, file_path: str, preview_path: str, page: int) -> Optional[str]:
        """Générer un aperçu d'une page PDF"""
        try:
            pages = convert_from_path(file_path, first_page=page, last_page=page, dpi=200)
            if pages:
                pages[0].save(preview_path, "PNG")
                return preview_path
        except Exception as e:
            logger.error(f"Erreur PDF preview: {str(e)}")
            return None
    
    def _generate_image_preview(self, file_path: str, preview_path: str) -> Optional[str]:
        """Générer un aperçu d'une image"""
        try:
            with Image.open(file_path) as img:
                # Redimensionner pour l'aperçu (max 800px de largeur)
                if img.width > 800:
                    ratio = 800 / img.width
                    new_height = int(img.height * ratio)
                    img = img.resize((800, new_height), Image.Resampling.LANCZOS)
                
                img.save(preview_path, "PNG")
                return preview_path
        except Exception as e:
            logger.error(f"Erreur image preview: {str(e)}")
            return None
    
    def perform_ocr(self, file_path: str, language: str = 'fra') -> Dict:
        """Effectuer la reconnaissance optique de caractères"""
        try:
            file_name = os.path.basename(file_path)
            name_without_ext = os.path.splitext(file_name)[0]
            ocr_result_path = os.path.join(self.ocr_folder, f"{name_without_ext}_ocr.json")
            
            file_ext = os.path.splitext(file_path)[1].lower()
            
            ocr_result = {
                'text': '',
                'confidence': 0,
                'language': language,
                'processed_at': datetime.now().isoformat(),
                'page_count': 0,
                'pages': []
            }
            
            if file_ext == '.pdf':
                ocr_result = self._perform_pdf_ocr(file_path, language)
            elif file_ext in ['.jpg', '.jpeg', '.png', '.tiff', '.bmp']:
                ocr_result = self._perform_image_ocr(file_path, language)
            
            # Sauvegarder le résultat
            with open(ocr_result_path, 'w', encoding='utf-8') as f:
                json.dump(ocr_result, f, ensure_ascii=False, indent=2)
            
            return ocr_result
            
        except Exception as e:
            logger.error(f"Erreur lors de l'OCR: {str(e)}")
            return {'error': str(e)}
    
    def _perform_pdf_ocr(self, file_path: str, language: str) -> Dict:
        """Effectuer l'OCR sur un PDF"""
        try:
            # Convertir PDF en images
            pages = convert_from_path(file_path, dpi=300)
            
            ocr_result = {
                'text': '',
                'confidence': 0,
                'language': language,
                'processed_at': datetime.now().isoformat(),
                'page_count': len(pages),
                'pages': []
            }
            
            total_confidence = 0
            full_text = ""
            
            for page_num, page in enumerate(pages, 1):
                # Convertir PIL Image en numpy array pour OpenCV
                open_cv_image = np.array(page)
                open_cv_image = cv2.cvtColor(open_cv_image, cv2.COLOR_RGB2BGR)
                
                # Préprocessing pour améliorer l'OCR
                processed_image = self._preprocess_image_for_ocr(open_cv_image)
                
                # Effectuer l'OCR
                data = pytesseract.image_to_data(processed_image, lang=language, output_type=pytesseract.Output.DICT)
                page_text = pytesseract.image_to_string(processed_image, lang=language)
                
                # Calculer la confiance moyenne
                confidences = [int(conf) for conf in data['conf'] if int(conf) > 0]
                page_confidence = sum(confidences) / len(confidences) if confidences else 0
                
                page_result = {
                    'page_number': page_num,
                    'text': page_text.strip(),
                    'confidence': page_confidence,
                    'word_count': len(page_text.split())
                }
                
                ocr_result['pages'].append(page_result)
                full_text += page_text + "\n"
                total_confidence += page_confidence
            
            ocr_result['text'] = full_text.strip()
            ocr_result['confidence'] = total_confidence / len(pages) if pages else 0
            
            return ocr_result
            
        except Exception as e:
            logger.error(f"Erreur PDF OCR: {str(e)}")
            return {'error': str(e)}
    
    def _perform_image_ocr(self, file_path: str, language: str) -> Dict:
        """Effectuer l'OCR sur une image"""
        try:
            # Charger l'image avec OpenCV
            image = cv2.imread(file_path)
            
            # Préprocessing
            processed_image = self._preprocess_image_for_ocr(image)
            
            # Effectuer l'OCR
            data = pytesseract.image_to_data(processed_image, lang=language, output_type=pytesseract.Output.DICT)
            text = pytesseract.image_to_string(processed_image, lang=language)
            
            # Calculer la confiance
            confidences = [int(conf) for conf in data['conf'] if int(conf) > 0]
            avg_confidence = sum(confidences) / len(confidences) if confidences else 0
            
            ocr_result = {
                'text': text.strip(),
                'confidence': avg_confidence,
                'language': language,
                'processed_at': datetime.now().isoformat(),
                'page_count': 1,
                'pages': [{
                    'page_number': 1,
                    'text': text.strip(),
                    'confidence': avg_confidence,
                    'word_count': len(text.split())
                }]
            }
            
            return ocr_result
            
        except Exception as e:
            logger.error(f"Erreur image OCR: {str(e)}")
            return {'error': str(e)}
    
    def _preprocess_image_for_ocr(self, image):
        """Préprocesser une image pour améliorer l'OCR"""
        try:
            # Convertir en niveaux de gris
            gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)
            
            # Appliquer un flou gaussien pour réduire le bruit
            blurred = cv2.GaussianBlur(gray, (5, 5), 0)
            
            # Seuillage adaptatif pour binariser l'image
            thresh = cv2.adaptiveThreshold(
                blurred, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C, cv2.THRESH_BINARY, 11, 2
            )
            
            # Morphologie pour nettoyer l'image
            kernel = np.ones((1, 1), np.uint8)
            processed = cv2.morphologyEx(thresh, cv2.MORPH_CLOSE, kernel)
            
            return processed
            
        except Exception as e:
            logger.error(f"Erreur preprocessing: {str(e)}")
            return image
    
    def update_document_metadata(self, document_id: int, metadata: Dict) -> bool:
        """Mettre à jour les métadonnées d'un document dans la base de données"""
        try:
            conn = db_connection()
            if not conn:
                return False
            
            cursor = conn.cursor()
            
            # Mettre à jour les métadonnées
            cursor.execute("""
                UPDATE document 
                SET metadonnees = %s, 
                    derniere_modification = NOW()
                WHERE id = %s
            """, (json.dumps(metadata), document_id))
            
            conn.commit()
            cursor.close()
            conn.close()
            
            return True
            
        except Exception as e:
            logger.error(f"Erreur mise à jour métadonnées: {str(e)}")
            return False
    
    def convert_document(self, file_path: str, target_format: str) -> Optional[str]:
        """Convertir un document vers un autre format"""
        try:
            file_name = os.path.basename(file_path)
            name_without_ext = os.path.splitext(file_name)[0]
            
            if target_format.lower() == 'pdf':
                output_path = os.path.join(self.upload_folder, f"{name_without_ext}_converted.pdf")
                return self._convert_to_pdf(file_path, output_path)
            elif target_format.lower() in ['jpg', 'png']:
                output_path = os.path.join(self.upload_folder, f"{name_without_ext}_converted.{target_format.lower()}")
                return self._convert_to_image(file_path, output_path, target_format)
            
            return None
            
        except Exception as e:
            logger.error(f"Erreur conversion: {str(e)}")
            return None
    
    def _convert_to_pdf(self, file_path: str, output_path: str) -> Optional[str]:
        """Convertir un document vers PDF"""
        # Implémentation basique - peut être étendue
        file_ext = os.path.splitext(file_path)[1].lower()
        
        if file_ext in ['.jpg', '.jpeg', '.png']:
            try:
                img = Image.open(file_path)
                if img.mode == 'RGBA':
                    img = img.convert('RGB')
                img.save(output_path, "PDF")
                return output_path
            except Exception as e:
                logger.error(f"Erreur conversion image vers PDF: {str(e)}")
                return None
        
        return None
    
    def _convert_to_image(self, file_path: str, output_path: str, format: str) -> Optional[str]:
        """Convertir un document vers une image"""
        file_ext = os.path.splitext(file_path)[1].lower()
        
        if file_ext == '.pdf':
            try:
                pages = convert_from_path(file_path, first_page=1, last_page=1, dpi=200)
                if pages:
                    pages[0].save(output_path, format.upper())
                    return output_path
            except Exception as e:
                logger.error(f"Erreur conversion PDF vers image: {str(e)}")
                return None
        
        return None